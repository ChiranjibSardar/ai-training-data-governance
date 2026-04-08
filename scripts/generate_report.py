"""Generate the RDI Framework POC Status Report as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import date


def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return heading


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


doc = Document()

# --- Title ---
title = doc.add_heading("RDI Framework — POC Status Report", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")
doc.add_paragraph("Version: 0.1.0 (Proof of Concept)")
doc.add_paragraph("Author: Chiranjib Sardar")
doc.add_paragraph("")

# --- Executive Summary ---
add_styled_heading(doc, "1. Executive Summary")
doc.add_paragraph(
    "This report summarizes the proof-of-concept (POC) implementation of the "
    "Responsible Data Infrastructure (RDI) Framework, a governance architecture "
    "for AI training data pipelines. The POC implements Layers 1 (Ingestion Gate) "
    "and Layer 2 (Validation & Risk Scoring) of the four-layer RDI architecture "
    "as a working Python package."
)
doc.add_paragraph(
    "The POC delivers four fully functional pipeline components — PII Scanner, "
    "Provenance Ledger, Toxicity Filter, and Deduplicator — along with a Risk "
    "Report generator, Pipeline orchestrator, and CLI interface. Four additional "
    "components (License Classifier, Quality Scorer, CMDI Calculator, C2PA Validator) "
    "are stubbed for Phase 2 implementation."
)
doc.add_paragraph(
    "All 67 unit tests pass with zero failures. The pipeline processes text files "
    "end-to-end: scanning for PII, scoring toxicity, detecting near-duplicates, "
    "logging to a hash-chained provenance ledger, and generating a structured "
    "JSON risk report. All tools, models, and datasets used are open-source."
)

# --- What Was Built ---
add_styled_heading(doc, "2. What Was Built")

add_styled_heading(doc, "2.1 Implemented Components (POC Scope)", level=2)
add_table(doc,
    ["Component", "Description", "Key Technology"],
    [
        ["PII Scanner", "Detects and redacts personally identifiable information (names, emails, phones, addresses, SSNs) from text", "Microsoft Presidio + spaCy en_core_web_lg"],
        ["Provenance Ledger", "Tamper-evident, SHA-256 hash-chained append-only audit log persisted as JSON Lines", "hashlib (Python stdlib)"],
        ["Toxicity Filter", "Multi-category toxicity scoring (toxic, severe_toxic, obscene, threat, insult, identity_hate)", "Detoxify (unitary/toxic-bert)"],
        ["Deduplicator", "Near-duplicate detection using MinHash with Locality-Sensitive Hashing", "datasketch"],
        ["Risk Report", "Structured JSON report with dataset summary, per-record results, and risk level assessment", "Python dataclasses + JSON"],
        ["Pipeline Orchestrator", "End-to-end pipeline coordinating all components across Layer 1 and Layer 2", "tqdm for progress"],
        ["CLI Interface", "Command-line interface with 'rdi run' and 'rdi validate-ledger' commands", "click"],
    ]
)

doc.add_paragraph("")

add_styled_heading(doc, "2.2 Stubbed Components (Phase 2)", level=2)
add_table(doc,
    ["Component", "Description", "Planned Technology"],
    [
        ["License Classifier", "ML-based classification of text documents by open-source license type", "HuggingFace Transformers (distilbert / bart-large-mnli)"],
        ["Quality Scorer", "Text quality scoring via perplexity measurement", "GPT-2 (HuggingFace)"],
        ["CMDI Calculator", "Cross-Modal Diversity Index measuring linguistic, topical, and geographic diversity", "langdetect + gensim LDA + spaCy NER"],
        ["C2PA Validator", "Content credential validation for media files", "c2patool CLI"],
    ]
)

doc.add_paragraph("")

add_styled_heading(doc, "2.3 Project Structure", level=2)
structure = doc.add_paragraph()
structure.style = doc.styles["No Spacing"]
lines = [
    "ai-training-data-governance/",
    "├── rdi/                          # Python package",
    "│   ├── __init__.py",
    "│   ├── models.py                 # 12 dataclasses with type hints",
    "│   ├── exceptions.py             # Exception hierarchy",
    "│   ├── pii_scanner.py            # PII detection & redaction",
    "│   ├── provenance_ledger.py      # Hash-chained audit log",
    "│   ├── toxicity_filter.py        # Multi-category toxicity scoring",
    "│   ├── deduplicator.py           # MinHash/LSH deduplication",
    "│   ├── risk_report.py            # Report generation & serialization",
    "│   ├── pipeline.py               # Pipeline orchestrator",
    "│   ├── cli.py                    # Click CLI interface",
    "│   ├── license_classifier.py     # [Stub] Phase 2",
    "│   ├── quality_scorer.py         # [Stub] Phase 2",
    "│   ├── cmdi_calculator.py        # [Stub] Phase 2",
    "│   └── c2pa_validator.py         # [Stub] Phase 2",
    "├── tests/                        # 67 unit tests",
    "│   ├── test_pii_scanner.py       # 9 tests",
    "│   ├── test_provenance_ledger.py # 17 tests",
    "│   ├── test_toxicity_filter.py   # 9 tests",
    "│   ├── test_deduplicator.py      # 9 tests",
    "│   ├── test_risk_report.py       # 12 tests",
    "│   ├── test_pipeline.py          # 6 tests",
    "│   └── test_cli.py              # 5 tests",
    "├── samples/                      # Sample data for demo",
    "├── demo.sh                       # Demo script",
    "└── pyproject.toml                # PEP 621 package config",
]
for line in lines:
    run = structure.add_run(line + "\n")
    run.font.name = "Courier New"
    run.font.size = Pt(8)

# --- Tests ---
add_styled_heading(doc, "3. Test Suite")

add_styled_heading(doc, "3.1 Test Summary", level=2)
add_table(doc,
    ["Test File", "Component", "Tests", "Passed", "Failed"],
    [
        ["test_pii_scanner.py", "PII Scanner", "9", "9", "0"],
        ["test_provenance_ledger.py", "Provenance Ledger", "17", "17", "0"],
        ["test_toxicity_filter.py", "Toxicity Filter", "9", "9", "0"],
        ["test_deduplicator.py", "Deduplicator", "9", "9", "0"],
        ["test_risk_report.py", "Risk Report", "12", "12", "0"],
        ["test_pipeline.py", "Pipeline Orchestrator", "6", "6", "0"],
        ["test_cli.py", "CLI Interface", "5", "5", "0"],
        ["TOTAL", "", "67", "67", "0"],
    ]
)

doc.add_paragraph("")

add_styled_heading(doc, "3.2 Test Categories", level=2)

doc.add_paragraph("PII Scanner Tests:", style="List Bullet")
pii_tests = [
    "Empty string input handling",
    "Email address detection and redaction ([EMAIL] placeholder)",
    "Person name detection and redaction ([PERSON] placeholder)",
    "Phone number detection and redaction ([PHONE] placeholder)",
    "SSN detection and redaction ([SSN] placeholder)",
    "Clean text (no PII) passes through unchanged",
    "Entity metadata validation (positions, confidence, original text)",
    "Entities sorted by start position",
]
for t in pii_tests:
    doc.add_paragraph(t, style="List Bullet 2")

doc.add_paragraph("Provenance Ledger Tests:", style="List Bullet")
ledger_tests = [
    "Genesis entry has previous_hash = '0' × 64",
    "Second entry chains to first entry's hash",
    "Entry hash is valid 64-character hex string",
    "All entry fields populated correctly",
    "Empty ledger verification succeeds",
    "Single and multi-entry verification succeeds",
    "Tampered entry_hash detected",
    "Tampered previous_hash detected",
    "Tampered content_hash detected",
    "JSONL file created on append",
    "Line count matches entry count",
    "Each line is valid JSON",
    "Write/read round-trip preserves data",
    "Reloaded ledger passes verification",
    "I/O error handling (malformed JSON, read-only directory)",
]
for t in ledger_tests:
    doc.add_paragraph(t, style="List Bullet 2")

doc.add_paragraph("Toxicity Filter Tests:", style="List Bullet")
tox_tests = [
    "Empty string returns all-zero scores",
    "Whitespace-only returns all-zero scores",
    "All six toxicity categories present in output",
    "All scores in [0.0, 1.0] range",
    "Benign text not flagged as high-risk",
    "is_high_risk matches threshold logic (any score > 0.8)",
    "Return type is ToxicityResult",
    "Benign greeting has low scores",
    "Benign factual text has low scores",
]
for t in tox_tests:
    doc.add_paragraph(t, style="List Bullet 2")

doc.add_paragraph("Deduplicator Tests:", style="List Bullet")
dedup_tests = [
    "Empty corpus returns empty results",
    "Identical documents placed in same cluster",
    "Identical documents have similarity pair",
    "Unrelated documents in separate clusters",
    "Every doc ID appears in exactly one cluster",
    "Similarity scores in [0.0, 1.0] range",
    "Lower threshold clusters more documents",
    "Custom num_perm still detects identical docs",
    "Return type is DeduplicationResult",
]
for t in dedup_tests:
    doc.add_paragraph(t, style="List Bullet 2")

doc.add_paragraph("Risk Report Tests:", style="List Bullet")
report_tests = [
    "generate_report produces valid RiskReport",
    "Risk level 'high' when records have high-risk toxicity",
    "Risk level 'high' when >50% flagged for review",
    "Risk level 'low' for clean data",
    "Risk level 'medium' when quality below threshold",
    "Risk level 'medium' when >20% duplicate ratio",
    "Empty records list handled correctly",
    "Dataset summary counts are accurate",
    "to_json produces valid JSON",
    "JSON round-trip preserves all fields",
    "Round-trip with similarity tuples",
    "Round-trip with empty records",
]
for t in report_tests:
    doc.add_paragraph(t, style="List Bullet 2")

# --- Datasets ---
add_styled_heading(doc, "4. Datasets Used")

doc.add_paragraph(
    "The POC uses a curated set of 5 sample text files designed to exercise "
    "each pipeline component. All sample data is synthetic — no real PII, "
    "proprietary data, or copyrighted content is used."
)

add_table(doc,
    ["File", "Purpose", "Content Description"],
    [
        ["clean_text.txt", "Baseline (no PII, non-toxic)", "Paragraph about responsible AI governance policies and data provenance"],
        ["pii_text.txt", "PII detection testing", "Text containing synthetic person names, email addresses, phone numbers, and SSN"],
        ["toxic_text.txt", "Toxicity scoring testing", "Mildly aggressive text with insults to trigger toxicity detection"],
        ["duplicate_a.txt", "Deduplication testing (original)", "Paragraph about machine learning, supervised/unsupervised learning, and deep learning"],
        ["duplicate_b.txt", "Deduplication testing (near-duplicate)", "Same paragraph as duplicate_a.txt with minor word substitutions (models→algorithms, make→generate, uses→utilizes)"],
    ]
)

doc.add_paragraph("")
doc.add_paragraph(
    "For unit testing, additional synthetic data is generated programmatically "
    "within each test file using pytest fixtures and temporary directories. "
    "No external datasets or pre-trained model fine-tuning datasets are required "
    "for the POC — all ML models (Presidio/spaCy, Detoxify) use their default "
    "pre-trained weights."
)

# --- Test Results ---
add_styled_heading(doc, "5. Summary of Test Results")

add_styled_heading(doc, "5.1 Overall Results", level=2)
add_table(doc,
    ["Metric", "Value"],
    [
        ["Total Tests", "67"],
        ["Passed", "67"],
        ["Failed", "0"],
        ["Pass Rate", "100%"],
        ["Execution Time", "~21 seconds"],
        ["Python Version", "3.12.5"],
        ["Test Framework", "pytest 9.0.3"],
        ["Platform", "macOS (darwin)"],
    ]
)

doc.add_paragraph("")

add_styled_heading(doc, "5.2 Component-Level Results", level=2)
doc.add_paragraph(
    "PII Scanner: All 9 tests pass. Successfully detects and redacts emails, "
    "person names, phone numbers, and SSNs. Clean text passes through unchanged. "
    "Entity metadata (positions, confidence scores, original text) is correctly populated."
)
doc.add_paragraph(
    "Provenance Ledger: All 17 tests pass. Hash chaining is correctly implemented "
    "with SHA-256. Tamper detection works for all corruption scenarios (modified "
    "entry hash, previous hash, and content hash). JSONL persistence round-trips "
    "without data loss. I/O error handling raises appropriate exceptions."
)
doc.add_paragraph(
    "Toxicity Filter: All 9 tests pass. Returns scores for all six toxicity "
    "categories. Benign text consistently scores below 0.5 across all categories. "
    "High-risk flagging logic correctly triggers when any score exceeds 0.8."
)
doc.add_paragraph(
    "Deduplicator: All 9 tests pass. Identical documents are reliably placed in "
    "the same cluster. Unrelated documents are separated. Every document ID appears "
    "in exactly one cluster. Configurable threshold affects clustering behavior as expected."
)
doc.add_paragraph(
    "Risk Report: All 12 tests pass. Risk level computation follows the defined "
    "thresholds (high/medium/low). JSON serialization and deserialization round-trip "
    "preserves all nested dataclass fields including PII entities and similarity tuples."
)
doc.add_paragraph(
    "Pipeline & CLI: All 11 tests pass. End-to-end pipeline processes single files "
    "and directories. CLI validates input paths, handles errors with appropriate exit "
    "codes, and produces correct output files."
)

# --- Metrics ---
add_styled_heading(doc, "6. Metrics")

add_styled_heading(doc, "6.1 Code Metrics", level=2)
add_table(doc,
    ["Metric", "Value"],
    [
        ["Source Files (rdi/)", "14"],
        ["Test Files (tests/)", "7"],
        ["Total Unit Tests", "67"],
        ["Data Model Classes", "12 dataclasses"],
        ["Exception Classes", "6"],
        ["CLI Commands", "2 (run, validate-ledger)"],
        ["Pipeline Components", "8 (4 implemented, 4 stubbed)"],
    ]
)

doc.add_paragraph("")

add_styled_heading(doc, "6.2 Open-Source Dependencies", level=2)
add_table(doc,
    ["Package", "Version", "Purpose", "License"],
    [
        ["presidio-analyzer", ">=2.2", "PII detection engine", "MIT"],
        ["presidio-anonymizer", ">=2.2", "PII redaction engine", "MIT"],
        ["spacy", ">=3.5", "NER model backend (en_core_web_lg)", "MIT"],
        ["detoxify", ">=0.5", "Toxicity classification", "Apache 2.0"],
        ["datasketch", ">=1.6", "MinHash/LSH deduplication", "MIT"],
        ["click", ">=8.0", "CLI framework", "BSD-3-Clause"],
        ["tqdm", ">=4.60", "Progress bars", "MIT/MPL"],
        ["hypothesis", ">=6.80", "Property-based testing (dev)", "MPL 2.0"],
        ["pytest", ">=7.0", "Test framework (dev)", "MIT"],
    ]
)

# --- Link to detailed test report ---
doc.add_paragraph("")
add_styled_heading(doc, "7. Detailed Test Report")
doc.add_paragraph(
    "The full detailed test output can be generated at any time by running:"
)
p = doc.add_paragraph()
run = p.add_run("python3 -m pytest tests/ -v --tb=long")
run.font.name = "Courier New"
run.font.size = Pt(10)

doc.add_paragraph(
    "For an HTML test report, install pytest-html and run:"
)
p = doc.add_paragraph()
run = p.add_run("pip install pytest-html\npython3 -m pytest tests/ -v --html=test_report.html --self-contained-html")
run.font.name = "Courier New"
run.font.size = Pt(10)

doc.add_paragraph(
    "The test report will be generated as test_report.html in the project root directory."
)

# --- How to Run ---
add_styled_heading(doc, "8. How to Run the Pipeline")

add_styled_heading(doc, "8.1 Prerequisites", level=2)
prereqs = [
    "Python 3.10 or higher",
    "pip (Python package manager)",
    "~2 GB disk space for ML model downloads (first run only)",
]
for p_text in prereqs:
    doc.add_paragraph(p_text, style="List Bullet")

add_styled_heading(doc, "8.2 Installation", level=2)
install_steps = [
    ("Clone the repository:", "git clone https://github.com/ChiranjibSardar/ai-training-data-governance.git\ncd ai-training-data-governance"),
    ("Install the package and dependencies:", "pip install -e '.[dev]'"),
    ("Download the spaCy language model:", "python -m spacy download en_core_web_lg"),
]
for desc, cmd in install_steps:
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

add_styled_heading(doc, "8.3 Running the Pipeline", level=2)

doc.add_paragraph("Option 1: Using the CLI")
p = doc.add_paragraph()
run = p.add_run("# Process a directory of text files\nrdi run --input /path/to/your/text/files/ --output /path/to/output/\n\n# Process a single file\nrdi run --input document.txt --output results/\n\n# With custom toxicity threshold\nrdi run --input data/ --output results/ --threshold-toxicity 0.9")
run.font.name = "Courier New"
run.font.size = Pt(10)

doc.add_paragraph("")
doc.add_paragraph("Option 2: Using the Python API")
p = doc.add_paragraph()
run = p.add_run(
    "from pathlib import Path\n"
    "from rdi.pipeline import Pipeline\n"
    "from rdi.models import PipelineConfig\n\n"
    "config = PipelineConfig(toxicity_threshold=0.8)\n"
    "pipeline = Pipeline(config=config)\n"
    "report = pipeline.run(\n"
    "    input_path=Path('your_data/'),\n"
    "    output_path=Path('output/'),\n"
    ")\n"
    "print(f'Risk level: {report.risk_level}')\n"
    "print(f'Records processed: {report.dataset_summary[\"total_records\"]}')"
)
run.font.name = "Courier New"
run.font.size = Pt(10)

doc.add_paragraph("")
doc.add_paragraph("Option 3: Using individual components")
p = doc.add_paragraph()
run = p.add_run(
    "from rdi.pii_scanner import PIIScanner\n"
    "from rdi.toxicity_filter import ToxicityFilter\n"
    "from rdi.deduplicator import Deduplicator\n\n"
    "# PII scanning\n"
    "scanner = PIIScanner()\n"
    "result = scanner.scan('Contact john@example.com')\n"
    "print(result.redacted_text)  # 'Contact [EMAIL]'\n\n"
    "# Toxicity scoring\n"
    "toxicity = ToxicityFilter()\n"
    "scores = toxicity.score('Some text to analyze')\n"
    "print(scores.scores)  # {'toxic': 0.01, ...}\n\n"
    "# Deduplication\n"
    "dedup = Deduplicator(threshold=0.8)\n"
    "result = dedup.deduplicate([('doc1', 'text...'), ('doc2', 'text...')])\n"
    "print(result.clusters)"
)
run.font.name = "Courier New"
run.font.size = Pt(10)

add_styled_heading(doc, "8.4 Validating the Provenance Ledger", level=2)
p = doc.add_paragraph()
run = p.add_run("rdi validate-ledger --ledger output/provenance_ledger.jsonl")
run.font.name = "Courier New"
run.font.size = Pt(10)

add_styled_heading(doc, "8.5 Running the Demo", level=2)
doc.add_paragraph("A demo script is included that processes the sample data:")
p = doc.add_paragraph()
run = p.add_run("chmod +x demo.sh\n./demo.sh")
run.font.name = "Courier New"
run.font.size = Pt(10)

add_styled_heading(doc, "8.6 Running Tests", level=2)
p = doc.add_paragraph()
run = p.add_run("# Run all tests\npython3 -m pytest tests/ -v\n\n# Run tests for a specific component\npython3 -m pytest tests/test_pii_scanner.py -v")
run.font.name = "Courier New"
run.font.size = Pt(10)

add_styled_heading(doc, "8.7 Input Data Requirements", level=2)
reqs = [
    "The pipeline processes .txt files (plain text, UTF-8 encoded)",
    "Input can be a single file or a directory containing multiple .txt files",
    "There is no minimum or maximum file size, but very large files may increase processing time",
    "The pipeline creates the output directory automatically if it does not exist",
    "Output includes: risk_report.json (structured report) and provenance_ledger.jsonl (audit log)",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

# --- Phase 2 Roadmap ---
add_styled_heading(doc, "9. Phase 2 Roadmap")
doc.add_paragraph(
    "The following components are planned for Phase 2 implementation:"
)
phase2 = [
    "License Classifier: Fine-tuned distilbert-base-uncased or zero-shot bart-large-mnli for classifying open-source license types (CC-BY, MIT, Apache, etc.)",
    "Quality Scorer: GPT-2 perplexity-based text quality scoring with sliding-window tokenization for documents up to 10,000 tokens",
    "CMDI Calculator: Cross-Modal Diversity Index computing linguistic (langdetect), topical (gensim LDA), and geographic (spaCy NER) diversity sub-indices",
    "C2PA Validator: Subprocess wrapper around c2patool CLI for validating content credentials on media files",
    "Property-Based Tests: Hypothesis-based property tests for all 17 correctness properties defined in the design document",
]
for item in phase2:
    doc.add_paragraph(item, style="List Bullet")

# Save
doc.save("reports/RDI_Framework_POC_Status_Report.docx")
print("Report saved to reports/RDI_Framework_POC_Status_Report.docx")
